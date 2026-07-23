#!/usr/bin/env python3
"""Export the authorized v2 prose manuscript to Markdown, DOCX, and PDF.

The v1 exporter pulled files from a remote repository and silently mixed local
overrides into them. This version is intentionally local and deterministic. It
also refuses to export prose until prose adaptation has been authorized.
"""

from __future__ import annotations

import argparse
import html
import json
import re
from dataclasses import dataclass
from pathlib import Path


REPO_ROOT = Path(__file__).resolve().parents[1]
APPROVAL_PATH = REPO_ROOT / "story" / "approval_status.json"
DEFAULT_SOURCE = REPO_ROOT / "story" / "timeline-weeks-prose-v2"
DEFAULT_OUTPUT = REPO_ROOT / "story" / "prose-v2-output"
EXPECTED_WEEKS = set(range(1, 55))
TITLE = "The Formula of Becoming"
SUBTITLE = "A Year at McCall-Hart University"


@dataclass(frozen=True)
class ProseFile:
    week: int
    path: Path
    content: str


def parse_week(path: Path) -> int | None:
    match = re.fullmatch(r"prose_(\d+)\.md", path.name, flags=re.IGNORECASE)
    return int(match.group(1)) if match else None


def require_prose_authorization() -> dict:
    approval = json.loads(APPROVAL_PATH.read_text(encoding="utf-8"))
    if not approval.get("prose_adaptation_allowed"):
        raise SystemExit(
            "Prose export is on hold: authorize prose adaptation first in "
            "story/approval_status.json."
        )
    return approval


def load_prose(source_dir: Path) -> list[ProseFile]:
    if not source_dir.is_dir():
        raise SystemExit(f"Prose source folder does not exist: {source_dir}")

    by_week: dict[int, ProseFile] = {}
    for path in sorted(source_dir.glob("*.md")):
        week = parse_week(path)
        if week is None:
            continue
        if week in by_week:
            raise SystemExit(f"Duplicate prose file for week {week}: {path}")
        by_week[week] = ProseFile(
            week=week,
            path=path,
            content=path.read_text(encoding="utf-8").strip(),
        )

    missing = sorted(EXPECTED_WEEKS - set(by_week))
    extras = sorted(set(by_week) - EXPECTED_WEEKS)
    if missing or extras:
        details = []
        if missing:
            details.append("missing " + ", ".join(str(week) for week in missing))
        if extras:
            details.append("unexpected " + ", ".join(str(week) for week in extras))
        raise SystemExit("Expected exactly weeks 1-54; " + "; ".join(details) + ".")
    return [by_week[week] for week in sorted(by_week)]


def markdown_parts(files: list[ProseFile]) -> list[str]:
    parts = [f"# {TITLE}", "", f"## {SUBTITLE}", ""]
    for index, prose in enumerate(files):
        if index:
            parts.extend(["", "---", ""])
        parts.append(prose.content)
    return parts


def write_markdown(files: list[ProseFile], destination: Path) -> None:
    destination.write_text("\n".join(markdown_parts(files)).rstrip() + "\n", encoding="utf-8")


def inline_markup(text: str) -> list[tuple[str, bool, bool]]:
    """Split the small Markdown emphasis subset used by the prose files."""
    pieces: list[tuple[str, bool, bool]] = []
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*|\*([^*\n]+?)\*", text):
        if match.start() > cursor:
            pieces.append((text[cursor:match.start()], False, False))
        if match.group(1) is not None:
            pieces.append((match.group(1), True, False))
        else:
            pieces.append((match.group(2), False, True))
        cursor = match.end()
    if cursor < len(text):
        pieces.append((text[cursor:], False, False))
    return pieces


def write_docx(files: list[ProseFile], destination: Path) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as exc:
        raise SystemExit("DOCX export requires python-docx.") from exc

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    section.different_first_page_header_footer = True

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.widow_control = True

    heading_tokens = {
        "Heading 1": (16, "2E74B5", 18, 10),
        "Heading 2": (13, "985424", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    header_run = header.add_run(TITLE.upper())
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(8.5)
    header_run.font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.space_after = Pt(0)
    page_run = footer.add_run()
    page_run.font.name = "Calibri"
    page_run.font.size = Pt(9)
    page_run.font.color.rgb = RGBColor(95, 95, 95)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    field_code = OxmlElement("w:instrText")
    field_code.set(qn("xml:space"), "preserve")
    field_code.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    page_run._r.extend([field_begin, field_code, field_end])

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(156)
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run(TITLE)
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(30)
    title_run.font.color.rgb = RGBColor(32, 55, 72)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(22)
    subtitle_run = subtitle.add_run(SUBTITLE)
    subtitle_run.italic = True
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(15)
    subtitle_run.font.color.rgb = RGBColor(152, 84, 36)

    edition = document.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edition_run = edition.add_run("Season One Prose Edition")
    edition_run.font.name = "Calibri"
    edition_run.font.size = Pt(10.5)
    edition_run.font.color.rgb = RGBColor(90, 90, 90)

    for prose in files:
        document.add_page_break()
        for line in prose.content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                paragraph = document.add_heading(stripped[4:], level=3)
            elif stripped.startswith("## "):
                paragraph = document.add_heading(stripped[3:], level=2)
            elif stripped.startswith("# "):
                paragraph = document.add_heading(stripped[2:], level=1)
            elif stripped.startswith(("- ", "* ")):
                paragraph = document.add_paragraph(style="List Bullet")
                stripped = stripped[2:]
            else:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.widow_control = True
            if not paragraph.text:
                for piece, bold, italic in inline_markup(stripped):
                    run = paragraph.add_run(piece)
                    run.bold = bold
                    run.italic = italic

    document.save(destination)


def reportlab_markup(text: str) -> str:
    marked_up = []
    for piece, bold, italic in inline_markup(text):
        escaped = html.escape(piece)
        if bold:
            escaped = f"<b>{escaped}</b>"
        elif italic:
            escaped = f"<i>{escaped}</i>"
        marked_up.append(escaped)
    return "".join(marked_up)


def write_pdf(files: list[ProseFile], destination: Path) -> None:
    try:
        from reportlab.lib.colors import HexColor
        from reportlab.lib.enums import TA_JUSTIFY
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise SystemExit("PDF export requires reportlab.") from exc

    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "TitleV2", parent=base["Title"], fontName="Helvetica-Bold",
            fontSize=28, leading=33, textColor=HexColor("#282D5A"), spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            "SubtitleV2", parent=base["Normal"], fontName="Helvetica-Oblique",
            fontSize=14, leading=18, alignment=1, textColor=HexColor("#985424"),
        ),
        "h1": ParagraphStyle(
            "H1V2", parent=base["Heading1"], fontName="Helvetica-Bold",
            fontSize=16, leading=20, textColor=HexColor("#2E74B5"),
            spaceBefore=18, spaceAfter=10, keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "H2V2", parent=base["Heading2"], fontName="Helvetica-Bold",
            fontSize=13, leading=17, textColor=HexColor("#985424"),
            spaceBefore=12, spaceAfter=6, keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "H3V2", parent=base["Heading3"], fontName="Helvetica-Bold",
            fontSize=12, leading=16, textColor=HexColor("#1F4D78"),
            spaceBefore=8, spaceAfter=4, keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "BodyV2", parent=base["BodyText"], fontName="Helvetica",
            fontSize=11, leading=14.66, alignment=TA_JUSTIFY, spaceAfter=8,
            allowWidows=False, allowOrphans=False,
        ),
        "bullet": ParagraphStyle(
            "BulletV2", parent=base["BodyText"], fontName="Helvetica",
            fontSize=10.5, leading=15, leftIndent=18, firstLineIndent=-10, spaceAfter=5,
        ),
    }

    flowables = [
        Spacer(1, 2.15 * inch),
        Paragraph(TITLE, styles["title"]),
        Paragraph(SUBTITLE, styles["subtitle"]),
        Spacer(1, 0.22 * inch),
        Paragraph("Season One Prose Edition", ParagraphStyle(
            "EditionV2", parent=base["Normal"], fontName="Helvetica",
            fontSize=10.5, leading=13, alignment=1, textColor=HexColor("#5A5A5A"),
        )),
    ]
    for prose in files:
        flowables.append(PageBreak())
        for line in prose.content.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("### "):
                flowables.append(Paragraph(reportlab_markup(stripped[4:]), styles["h3"]))
            elif stripped.startswith("## "):
                flowables.append(Paragraph(reportlab_markup(stripped[3:]), styles["h2"]))
            elif stripped.startswith("# "):
                flowables.append(Paragraph(reportlab_markup(stripped[2:]), styles["h1"]))
            elif stripped.startswith(("- ", "* ")):
                flowables.append(Paragraph("&bull; " + reportlab_markup(stripped[2:]), styles["bullet"]))
            else:
                flowables.append(Paragraph(reportlab_markup(stripped), styles["body"]))

    document = SimpleDocTemplate(
        str(destination), pagesize=letter, topMargin=0.85 * inch,
        bottomMargin=0.85 * inch, leftMargin=0.9 * inch, rightMargin=0.9 * inch,
        title=TITLE, author="McCall-Hart University Comic Project",
    )

    def first_page(canvas, doc):
        canvas.saveState()
        canvas.setTitle(TITLE)
        canvas.restoreState()

    def later_pages(canvas, doc):
        canvas.saveState()
        width, _ = letter
        canvas.setFillColor(HexColor("#646464"))
        canvas.setFont("Helvetica", 8.5)
        canvas.drawCentredString(width / 2, 10.35 * inch, TITLE.upper())
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(width / 2, 0.42 * inch, str(doc.page))
        canvas.restoreState()

    document.build(flowables, onFirstPage=first_page, onLaterPages=later_pages)


def word_count(files: list[ProseFile]) -> int:
    return sum(len(re.findall(r"\b[\w'-]+\b", prose.content)) for prose in files)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--source-dir", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output-dir", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    approval = require_prose_authorization()
    files = load_prose(args.source_dir.resolve())
    output_dir = args.output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)

    basename = "the-formula-of-becoming-prose-v2"
    write_markdown(files, output_dir / f"{basename}.md")
    write_docx(files, output_dir / f"{basename}.docx")
    write_pdf(files, output_dir / f"{basename}.pdf")

    summary = {
        "setting": approval.get("setting"),
        "script_version": approval.get("script_version"),
        "week_count": len(files),
        "word_count": word_count(files),
        "source_dir": str(args.source_dir.resolve().relative_to(REPO_ROOT)),
        "output_dir": str(output_dir.relative_to(REPO_ROOT)),
    }
    (output_dir / "export-summary.json").write_text(
        json.dumps(summary, indent=2) + "\n", encoding="utf-8"
    )
    print(
        f"Exported {summary['week_count']} prose files ({summary['word_count']:,} words) "
        f"to {summary['output_dir']}."
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
